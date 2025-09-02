import asyncio
import io
import traceback
import logging
from docx import Document

logger = logging.getLogger(__name__)

# Import services with error handling
try:
    from services.protocol_parser import extract_drugs_from_text, generate_document_summary
    from services.regulatory_checker import check_all_regulators
    from services.pubmed_client import PubMedClient
    logger.info("✅ All services imported successfully")
except ImportError as e:
    logger.error(f"❌ Service import error: {e}")
    raise

# Create PubMed client
pubmed = PubMedClient()

async def process_single_drug(drug_info: dict, document_context: str = ""):
    """Process a single drug with regulatory and PubMed checks"""
    source_name = drug_info.get("drug_name_source")
    inn_name = drug_info.get("inn_name", "")
    
    if not source_name:
        logger.warning(f"❌ Skipping drug with no source name: {drug_info}")
        return None

    logger.info(f"💊 Processing drug: {source_name} (INN: {inn_name})")
    
    # If no INN found by Gemini, return basic data
    if not inn_name or inn_name.lower() in ['unknown', 'не определен', 'не определено', '']:
        logger.warning(f"⚠️ No INN found for {source_name}, returning basic data")
        return {
            "source_data": drug_info,
            "normalization": {"inn_name": None, "source": "Gemini", "confidence": "none"},
            "regulatory_checks": {"regulatory_checks": {}, "dosage_check": {}},
            "pubmed_articles": [],
            "ai_analysis": {
                "ud_ai_grade": drug_info.get("ud_ai_grade", "Unknown"),
                "ud_ai_justification": "Не удалось определить МНН препарата",
                "ai_summary_note": drug_info.get("ai_note", "Требуется ручная проверка названия препарата")
            }
        }

    logger.info(f"✅ Found INN: {inn_name}, proceeding with checks")
    
    try:
        # Enhanced context for PubMed search
        search_context = drug_info.get("context_indication", "")
        if not search_context and document_context:
            search_context = document_context[:200]
        
        # Run regulatory and PubMed checks in parallel
        logger.info("🔄 Running regulatory and PubMed tasks in parallel...")
        regulatory_task = check_all_regulators(
            inn_name=inn_name,
            source_dosage=drug_info.get("dosage_source", "")
        )
        
        pubmed_task = pubmed.search_articles_for_drug(
            inn_name=inn_name,
            brand_name=source_name,
            context=search_context
        )

        regulatory_results, pubmed_articles = await asyncio.gather(
            regulatory_task, 
            pubmed_task, 
            return_exceptions=True
        )
        
        # Handle exceptions in results
        if isinstance(regulatory_results, Exception):
            logger.error(f"❌ Regulatory check failed: {regulatory_results}")
            regulatory_results = {"regulatory_checks": {}, "dosage_check": {}}
            
        if isinstance(pubmed_articles, Exception):
            logger.error(f"❌ PubMed search failed: {pubmed_articles}")
            pubmed_articles = []
            
    except Exception as e:
        logger.error(f"❌ Error in parallel tasks: {e}")
        regulatory_results = {"regulatory_checks": {}, "dosage_check": {}}
        pubmed_articles = []

    # Build full drug data
    full_drug_data = {
        "source_data": drug_info,
        "normalization": {"inn_name": inn_name, "source": "Gemini", "confidence": "high"},
        "regulatory_checks": regulatory_results,
        "pubmed_articles": pubmed_articles,
        "ai_analysis": {
            "ud_ai_grade": drug_info.get("ud_ai_grade", "Unknown"),
            "ud_ai_justification": drug_info.get("formulary_status", "Анализ не выполнен"),
            "ai_summary_note": drug_info.get("ai_note", "Заметка не сгенерирована")
        }
    }
    
    logger.info(f"✅ Completed analysis for: {source_name}")
    return full_drug_data

async def run_analysis_pipeline(file_content: bytes):
    """Main analysis pipeline using demo approach"""
    logger.info("🚀 Starting DEMO-STYLE analysis pipeline")
    logger.info(f"File content size: {len(file_content)} bytes")
    
    file_stream = io.BytesIO(file_content)
    
    try:
        logger.info("📄 Loading DOCX document...")
        document = Document(file_stream)
        logger.info("✅ Document loaded successfully")
    except Exception as e:
        logger.error(f"❌ Error loading document: {e}")
        return {"error": f"Не удалось загрузить документ: {e}"}

    logger.info("📄 Extracting text from document...")
    all_text = []
    
    # Extract from paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            all_text.append(paragraph.text)
    
    # Extract from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text)
    
    full_text = "\n".join(all_text)
    logger.info(f"📄 Extracted text length: {len(full_text)} chars")

    if not full_text.strip():
        logger.error("❌ Document is empty")
        return {"error": "Документ пуст или не содержит текста."}

    try:
        # Generate summary and extract drugs concurrently
        logger.info("🔄 Starting summary and drug extraction")
        summary_task = generate_document_summary(full_text)
        extraction_task = extract_drugs_from_text(full_text)

        logger.info("⏳ Waiting for summary and extraction tasks...")
        document_summary, extracted_drugs = await asyncio.gather(
            summary_task, 
            extraction_task,
            return_exceptions=True
        )
        
        # Handle exceptions
        if isinstance(document_summary, Exception):
            logger.error(f"❌ Summary generation failed: {document_summary}")
            document_summary = "Не удалось сгенерировать резюме"
            
        if isinstance(extracted_drugs, Exception):
            logger.error(f"❌ Drug extraction failed: {extracted_drugs}")
            extracted_drugs = []
        
        logger.info(f"✅ Summary and extraction completed. Found {len(extracted_drugs)} drugs")
        
    except Exception as e:
        logger.error(f"❌ Error in summary/extraction: {e}")
        return {"error": f"Ошибка при анализе документа: {e}"}

    if not extracted_drugs:
        logger.warning("⚠️ No drugs found in document")
        return {
            "document_summary": document_summary,
            "analysis_results": [],
            "message": "В документе не найдено лекарственных препаратов."
        }

    logger.info(f"🔄 Processing {len(extracted_drugs)} extracted drugs")
    for i, drug in enumerate(extracted_drugs):
        logger.info(f"  Drug {i+1}: {drug.get('drug_name_source', 'Unknown')} (INN: {drug.get('inn_name', 'Unknown')})")

    try:
        logger.info("🔄 Starting parallel drug enhancement...")
        analysis_tasks = [process_single_drug(drug, full_text[:1000]) for drug in extracted_drugs]
        analysis_results = await asyncio.gather(*analysis_tasks, return_exceptions=True)
        logger.info("✅ All drug enhancement tasks completed")
        
        # Filter out exceptions and None results
        valid_results = []
        for i, result in enumerate(analysis_results):
            if isinstance(result, Exception):
                logger.error(f"❌ Enhancement failed for drug {i}: {result}")
            elif result is not None:
                valid_results.append(result)
                logger.info(f"✅ Drug {i+1} enhancement successful")
        
        logger.info(f"✅ Pipeline completed. {len(valid_results)} drugs analyzed successfully")
        
        return {
            "document_summary": document_summary,
            "analysis_results": valid_results
        }
        
    except Exception as e:
        logger.error(f"❌ Error in drug enhancement: {e}")
        logger.error(traceback.format_exc())
        return {"error": f"Ошибка при анализе препаратов: {e}"}
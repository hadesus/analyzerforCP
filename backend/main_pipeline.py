import asyncio
import io
import traceback
import logging
from docx import Document

logger = logging.getLogger(__name__)

# Import all our services
logger.info("Importing services...")
try:
    logger.info("Importing protocol_parser...")
    from services import protocol_parser
    logger.info("✅ protocol_parser imported")
    
    logger.info("Importing drug_normalizer...")
    from services import drug_normalizer
    logger.info("✅ drug_normalizer imported")
    
    logger.info("Importing regulatory_checker...")
    from services import regulatory_checker
    logger.info("✅ regulatory_checker imported")
    
    logger.info("Importing pubmed_client...")
    from services import pubmed_client
    logger.info("✅ pubmed_client imported")
    
    logger.info("Importing ai_analyzer...")
    from services import ai_analyzer
    logger.info("✅ ai_analyzer imported")
    
except ImportError as e:
    logger.error(f"❌ Import error: {e}")
    logger.error(f"Traceback: {traceback.format_exc()}")
    traceback.print_exc()
    raise

# Instantiate clients/services
logger.info("Creating PubMed client...")
pubmed = pubmed_client.PubMedClient()
logger.info("✅ PubMed client created")

async def generate_document_summary(full_text: str) -> str:
    """Uses Gemini to generate a brief summary of the entire document."""
    logger.info("📝 Starting document summary generation...")
    if not full_text:
        logger.warning("❌ Empty text for summary generation")
        return "Текст документа пуст."

    logger.info(f"📝 Generating summary for text of length: {len(full_text)}")
    
    prompt = f"""Проанализируй клинический протокол и напиши краткое резюме (2-3 предложения).

Укажи:
- Основное заболевание или состояние
- Целевую группу пациентов
- Основные подходы к лечению

Текст протокола:
---
{full_text[:15000]}
---

Резюме:"""
    
    try:
        logger.info("🤖 Sending summary request to Gemini...")
        import google.generativeai as genai
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            generation_config={"temperature": 0.2, "max_output_tokens": 500}
        )
        response = await model.generate_content_async(prompt)
        summary = response.text.strip()
        logger.info(f"✅ Summary generated: {len(summary)} chars")
        return summary
    except Exception as e:
        logger.error(f"❌ Error generating document summary: {e}")
        return "Не удалось сгенерировать резюме документа."

async def process_single_drug(drug_info: dict, document_context: str = ""):
    """
    Processes a single drug through the complete analysis pipeline.
    Enhanced to use existing INN from extraction and better context.
    """
    source_name = drug_info.get("drug_name_source")
    existing_inn = drug_info.get("inn_name", "")
    
    if not source_name:
        logger.warning(f"❌ Skipping drug with no source name: {drug_info}")
        return None

    logger.info(f"💊 Processing drug: {source_name}")
    
    try:
        logger.info(f"🔄 Normalizing drug name: {source_name}")
        # Use existing INN if provided by extraction
        if existing_inn and existing_inn.lower() not in ['unknown', 'не определен']:
            normalization_result = await drug_normalizer.normalize_drug_with_existing_inn(source_name, existing_inn)
        else:
            normalization_result = await drug_normalizer.normalize_drug(source_name)
        
        inn_name = normalization_result.get("inn_name")
        logger.info(f"✅ Normalization result: {normalization_result}")
    except Exception as e:
        logger.error(f"❌ Normalization failed for {source_name}: {e}")
        normalization_result = {"inn_name": None, "source": "Error", "confidence": "none"}
        inn_name = None

    if not inn_name:
        logger.warning(f"⚠️ No INN found for {source_name}, returning partial data")
        return {
            "source_data": drug_info,
            "normalization": normalization_result,
            "regulatory_checks": {"regulatory_checks": {}, "dosage_check": {}},
            "pubmed_articles": [],
            "ai_analysis": {
                "ud_ai_grade": "Unknown",
                "ud_ai_justification": "Не удалось определить МНН препарата",
                "ai_summary_note": "Требуется ручная проверка названия препарата"
            }
        }

    logger.info(f"✅ Found INN: {inn_name}, proceeding with full analysis")
    
    try:
        logger.info(f"🔍 Starting regulatory and PubMed checks for {inn_name}")
        
        # Enhanced context for PubMed search
        search_context = drug_info.get("context_indication", "")
        if not search_context and document_context:
            # Use document context if drug context is empty
            search_context = document_context[:200]  # First 200 chars of document
        
        # Run regulatory and PubMed checks in parallel
        regulatory_task = regulatory_checker.check_all_regulators(
            inn_name=inn_name,
            source_dosage=drug_info.get("dosage_source", "")
        )
        
        pubmed_task = pubmed.search_articles(
            inn_name=inn_name,
            brand_name=source_name,
            disease_context=search_context
        )

        logger.info("🔄 Running regulatory and PubMed tasks in parallel...")
        regulatory_results, pubmed_articles = await asyncio.gather(
            regulatory_task, 
            pubmed_task, 
            return_exceptions=True
        )
        logger.info("✅ Parallel tasks completed")
        
        # Handle exceptions in results
        if isinstance(regulatory_results, Exception):
            logger.error(f"❌ Regulatory check failed: {regulatory_results}")
            regulatory_results = {"regulatory_checks": {}, "dosage_check": {}}
            
        if isinstance(pubmed_articles, Exception):
            logger.error(f"❌ PubMed search failed: {pubmed_articles}")
            pubmed_articles = []
            
    except Exception as e:
        logger.error(f"❌ Error in parallel tasks: {e}")
        regulatory_results = {"regulatory_checks": {}, "dosage_check": {}}
        pubmed_articles = []

    # Build full drug data
    full_drug_data = {
        "source_data": drug_info,
        "normalization": normalization_result,
        "regulatory_checks": regulatory_results,
        "pubmed_articles": pubmed_articles,
    }

    try:
        logger.info(f"🧠 Generating final analysis for {source_name}")
        final_analysis = await ai_analyzer.generate_ai_analysis(full_drug_data)
        logger.info(f"✅ Analysis completed for {source_name}")
    except Exception as e:
        logger.error(f"❌ Analysis failed for {source_name}: {e}")
        final_analysis = {
            "ud_ai_grade": "Error",
            "ud_ai_justification": "Ошибка при генерации анализа",
            "ai_summary_note": "Не удалось сгенерировать анализ"
        }
        
    full_drug_data["ai_analysis"] = final_analysis
    
    logger.info(f"✅ Completed full analysis for: {source_name}")
    return full_drug_data

async def run_analysis_pipeline(file_content: bytes):
    """
    Enhanced main orchestrator using unified extraction approach like demo.
    """
    logger.info("🚀 Starting enhanced analysis pipeline")
    logger.info(f"File content size: {len(file_content)} bytes")
    
    file_stream = io.BytesIO(file_content)
    
    try:
        logger.info("📄 Loading DOCX document...")
        document = Document(file_stream)
        logger.info("✅ Document loaded successfully")
    except Exception as e:
        logger.error(f"❌ Error loading document: {e}")
        return {"error": f"Не удалось загрузить документ: {e}"}

    logger.info("📄 Extracting text from document...")
    all_text = []
    
    # Extract from paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            all_text.append(paragraph.text)
    
    # Extract from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text)
    
    full_text = "\n".join(all_text)
    logger.info(f"📄 Extracted text length: {len(full_text)} chars")

    if not full_text.strip():
        logger.error("❌ Document is empty")
        return {"error": "Документ пуст или не содержит текста."}

    try:
        # Generate summary and extract drugs concurrently
        logger.info("🔄 Starting summary and unified drug extraction")
        summary_task = generate_document_summary(full_text)
        extraction_task = protocol_parser.extract_drugs_from_text(full_text)

        logger.info("⏳ Waiting for summary and extraction tasks...")
        document_summary, extracted_drugs = await asyncio.gather(
            summary_task, 
            extraction_task,
            return_exceptions=True
        )
        
        # Handle exceptions
        if isinstance(document_summary, Exception):
            logger.error(f"❌ Summary generation failed: {document_summary}")
            document_summary = "Не удалось сгенерировать резюме"
            
        if isinstance(extracted_drugs, Exception):
            logger.error(f"❌ Drug extraction failed: {extracted_drugs}")
            extracted_drugs = []
        
        logger.info(f"✅ Summary and extraction completed. Found {len(extracted_drugs)} drugs")
        
    except Exception as e:
        logger.error(f"❌ Error in summary/extraction: {e}")
        return {"error": f"Ошибка при анализе документа: {e}"}

    if not extracted_drugs:
        logger.warning("⚠️ No drugs found in document")
        return {
            "document_summary": document_summary,
            "analysis_results": [],
            "message": "В документе не найдено лекарственных препаратов."
        }

    logger.info(f"🔄 Processing {len(extracted_drugs)} extracted drugs")
    for i, drug in enumerate(extracted_drugs):
        logger.info(f"  Drug {i+1}: {drug.get('drug_name_source', 'Unknown')} (INN: {drug.get('inn_name', 'Unknown')})")

    try:
        logger.info("🔄 Starting parallel drug analysis...")
        # Pass document context to each drug processing
        analysis_tasks = [process_single_drug(drug, full_text[:1000]) for drug in extracted_drugs]
        analysis_results = await asyncio.gather(*analysis_tasks, return_exceptions=True)
        logger.info("✅ All drug analysis tasks completed")
        
        # Filter out exceptions and None results
        valid_results = []
        for i, result in enumerate(analysis_results):
            if isinstance(result, Exception):
                logger.error(f"❌ Analysis failed for drug {i}: {result}")
            elif result is not None:
                valid_results.append(result)
                logger.info(f"✅ Drug {i+1} analysis successful")
        
        logger.info(f"✅ Pipeline completed. {len(valid_results)} drugs analyzed successfully")
        
        return {
            "document_summary": document_summary,
            "analysis_results": valid_results
        }
        
    except Exception as e:
        logger.error(f"❌ Error in drug analysis: {e}")
        logger.error(traceback.format_exc())
        return {"error": f"Ошибка при анализе препаратов: {e}"}
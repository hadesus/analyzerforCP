import io
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook

def create_docx_export(data: list) -> io.BytesIO:
    """Creates a .docx file from the analysis data and returns it as a byte stream."""
    document = Document()
    document.add_heading('Анализ Клинического Протокола', 0)

    # Define table headers
    headers = [
        'Название (Источник)', 'МНН', 'Дозировка (Источник)',
        'Уровень док-ти (AI/GRADE)', 'Обоснование (AI)', 'Заметка (AI)',
        'Статус FDA', 'Статус EMA', 'Статус BNF', 'Статус WHO EML',
        'Сравнение доз', 'PubMed Ссылки'
    ]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for item in data:
        row_cells = table.add_row().cells
        # A helper to safely get nested data
        def get_val(path, default='N/A'):
            keys = path.split('.')
            val = item
            for key in keys:
                val = val.get(key) if isinstance(val, dict) else None
                if val is None:
                    return default
            return val

        row_cells[0].text = get_val('source_data.drug_name_source')
        row_cells[1].text = get_val('normalization.inn_name')
        row_cells[2].text = get_val('source_data.dosage_source')
        row_cells[3].text = get_val('ai_analysis.ud_ai_grade')
        row_cells[4].text = get_val('ai_analysis.ud_ai_justification')
        row_cells[5].text = get_val('ai_analysis.ai_summary_note')
        row_cells[6].text = get_val('regulatory_checks.FDA.status')
        row_cells[7].text = get_val('regulatory_checks.EMA.status')
        row_cells[8].text = get_val('regulatory_checks.BNF.status')
        row_cells[9].text = get_val('regulatory_checks.WHO_EML.status')
        row_cells[10].text = get_val('dosage_check.comparison_result')

        # Handle links
        links = get_val('pubmed_articles', [])
        if links:
            row_cells[11].text = '\n'.join([f"{a.get('title', '')} ({a.get('link', '')})" for a in links])
        else:
            row_cells[11].text = 'N/A'


    # Save to a byte stream
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_excel_export(data: list) -> io.BytesIO:
    """Creates an .xlsx file from the analysis data."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Анализ Протокола"

    headers = [
        'Название (Источник)', 'МНН', 'Дозировка (Источник)',
        'Уровень док-ти (AI/GRADE)', 'Обоснование (AI)', 'Заметка (AI)',
        'Статус FDA', 'Статус EMA', 'Статус BNF', 'Статус WHO EML',
        'Сравнение доз', 'PubMed Ссылки'
    ]
    ws.append(headers)

    for item in data:
        def get_val(path, default='N/A'):
            keys = path.split('.')
            val = item
            for key in keys:
                val = val.get(key) if isinstance(val, dict) else None
                if val is None:
                    return default
            return val

        links = get_val('pubmed_articles', [])
        link_text = '\n'.join([a.get('link', '') for a in links]) if links else 'N/A'

        row = [
            get_val('source_data.drug_name_source'),
            get_val('normalization.inn_name'),
            get_val('source_data.dosage_source'),
            get_val('ai_analysis.ud_ai_grade'),
            get_val('ai_analysis.ud_ai_justification'),
            get_val('ai_analysis.ai_summary_note'),
            get_val('regulatory_checks.FDA.status'),
            get_val('regulatory_checks.EMA.status'),
            get_val('regulatory_checks.BNF.status'),
            get_val('regulatory_checks.WHO_EML.status'),
            get_val('dosage_check.comparison_result'),
            link_text
        ]
        ws.append(row)

    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    return file_stream
